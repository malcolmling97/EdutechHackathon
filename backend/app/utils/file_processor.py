"""
File processing utilities for text extraction.

Handles text extraction from various file formats:
- PDF: Using PyPDF2 or pdfplumber
- DOCX: Using python-docx
- TXT/MD: Direct reading

Based on BACKEND_DATA_FLOWS.md section 5.1 - Upload Files Flow.
"""
import os
import uuid
import magic
import aiofiles
from pathlib import Path
from typing import Optional, Tuple
from fastapi import UploadFile
import logging

# Text extraction dependencies (will install these)
try:
    import PyPDF2
    from pdfplumber import PDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.core.config import settings

logger = logging.getLogger(__name__)


class FileProcessor:
    """File processing utility for text extraction and file handling."""
    
    def __init__(self, upload_dir: str = None):
        """Initialize file processor with upload directory."""
        self.upload_dir = Path(upload_dir or settings.upload_dir)
        self.temp_dir = Path(settings.temp_dir)
        
        # Ensure directories exist
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.temp_dir.mkdir(parents=True, exist_ok=True)
    
    async def save_uploaded_file(self, file: UploadFile) -> Tuple[str, str]:
        """
        Save uploaded file to storage and return path and unique filename.
        
        Args:
            file: FastAPI UploadFile instance
            
        Returns:
            Tuple of (storage_path, unique_filename)
        """
        # Generate unique filename to avoid conflicts
        file_extension = self._get_file_extension(file.filename)
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        storage_path = self.upload_dir / unique_filename
        
        # Save file using async file operations
        async with aiofiles.open(storage_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        logger.info(f"Saved file {file.filename} as {unique_filename}")
        return str(storage_path), unique_filename
    
    def detect_mime_type(self, file_path: str) -> str:
        """
        Detect MIME type of file using python-magic.
        
        Args:
            file_path: Path to the file
            
        Returns:
            MIME type string
        """
        try:
            mime = magic.Magic(mime=True)
            return mime.from_file(file_path)
        except Exception as e:
            logger.warning(f"Could not detect MIME type for {file_path}: {e}")
            # Fallback to extension-based detection
            return self._mime_from_extension(file_path)
    
    def _mime_from_extension(self, file_path: str) -> str:
        """Fallback MIME type detection based on file extension."""
        extension = self._get_file_extension(file_path).lower()
        
        mime_map = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.txt': 'text/plain',
            '.md': 'text/markdown'
        }
        
        return mime_map.get(extension, 'application/octet-stream')
    
    def _get_file_extension(self, filename: str) -> str:
        """Extract file extension from filename."""
        return Path(filename).suffix if filename else ''
    
    async def extract_text_content(self, file_path: str, mime_type: str) -> Optional[str]:
        """
        Extract text content from file based on MIME type.
        
        Args:
            file_path: Path to the file
            mime_type: MIME type of the file
            
        Returns:
            Extracted text content or None if extraction failed
        """
        try:
            if mime_type == 'application/pdf':
                return await self._extract_pdf_text(file_path)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return await self._extract_docx_text(file_path)
            elif mime_type in ['text/plain', 'text/markdown']:
                return await self._extract_text_file(file_path)
            else:
                logger.warning(f"Unsupported MIME type for text extraction: {mime_type}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None
    
    async def _extract_pdf_text(self, file_path: str) -> Optional[str]:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            logger.error("PDF processing libraries not available")
            return None
        
        try:
            # Try with pdfplumber first (better text extraction)
            try:
                with PDF.open(file_path) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    return '\n'.join(text_parts)
            except:
                # Fallback to PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_parts = []
                    
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    
                    return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return None
    
    async def _extract_docx_text(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            logger.error("DOCX processing library not available")
            return None
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return None
    
    async def _extract_text_file(self, file_path: str) -> Optional[str]:
        """Extract text from plain text or markdown file."""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                return await file.read()
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    async with aiofiles.open(file_path, 'r', encoding=encoding) as file:
                        content = await file.read()
                        logger.info(f"Successfully read file with {encoding} encoding")
                        return content
                except UnicodeDecodeError:
                    continue
            
            logger.error(f"Could not decode text file {file_path}")
            return None
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return None
    
    def delete_file(self, file_path: str) -> bool:
        """
        Delete file from storage.
        
        Args:
            file_path: Path to the file to delete
            
        Returns:
            True if successful, False otherwise
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Deleted file: {file_path}")
                return True
            else:
                logger.warning(f"File not found: {file_path}")
                return False
        except Exception as e:
            logger.error(f"Error deleting file {file_path}: {e}")
            return False
    
    def get_file_size(self, file_path: str) -> int:
        """Get file size in bytes."""
        try:
            return os.path.getsize(file_path)
        except OSError:
            return 0


# Global file processor instance
file_processor = FileProcessor() 